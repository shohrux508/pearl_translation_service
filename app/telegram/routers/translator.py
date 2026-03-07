import os
import asyncio
from pathlib import Path
from docx import Document
from aiogram import Router, types, F
from aiogram.types import FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State

from app.container import Container
from app.services.document_manager import doc_manager

router = Router()
container_instance: Container = None

class TranslationState(StatesGroup):
    waiting_for_back_photo = State()
    choosing_doc_type = State()
    choosing_language = State()
    validating_data = State()
    editing_field = State()

def setup_router(container: Container):
    """Инжектим контейнер в роутер."""
    global container_instance
    container_instance = container

def get_gemini_service():
    """Безопасное получение сервиса Gemini."""
    try:
        return container_instance.get("gemini_service")
    except (KeyError, AttributeError):
        return None

def _create_temp_template(template_path: Path, doc_name: str, lang_name: str):
    """Синхронная функция создания пустого шаблона (выполнять в to_thread)."""
    if not template_path.exists():
        template_path.parent.mkdir(exist_ok=True, parents=True)
        doc = Document()
        doc.add_paragraph(f"Временный тестовый шаблон для: {doc_name}")
        doc.add_paragraph(f"Язык шаблона (и перевода): {lang_name}")
        doc.add_paragraph("Данные будут автоматически вставляться, если шаблонизатор найдет совпадающие поля (например {{surname}}).")
        doc.save(template_path)

def _generate_docx(gemini_service, extracted_data: dict, template_path: Path, output_path: Path):
    """Синхронная генерация Word-документа."""
    gemini_service.insert_into_docx(
        data=extracted_data,
        template_path=template_path,
        output_path=output_path
    )

@router.message(F.photo)
async def handle_document_photo(message: types.Message, state: FSMContext):
    """Принимает первую или вторую фотографию документа."""
    gemini_service = get_gemini_service()
    if not gemini_service:
        await message.reply("❌ Сервис перевода (Gemini) в данный момент недоступен. Проверьте API ключ.")
        return

    current_state = await state.get_state()
    if current_state == TranslationState.waiting_for_back_photo.state:
        # Это второе фото (задняя сторона)
        photo = message.photo[-1]
        data = await state.get_data()
        file_ids = data.get("file_ids", [])
        file_ids.append(photo.file_id)
        await state.update_data(file_ids=file_ids)
        await ask_for_doc_type(message, state)
        return

    # Иначе это первое фото: стартуем новую цепочку
    await state.clear()
    photo = message.photo[-1]
    await state.update_data(file_ids=[photo.file_id])
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➡️ Без задней части (пропустить)", callback_data="skip_back_photo")]
    ])
    
    await message.reply(
        "📸 Передняя часть документа получена!\n\nЕсли у документа есть задняя часть, отправьте её сейчас (другой фотографией).\nЕсли задней части нет, нажмите кнопку ниже:",
        reply_markup=keyboard
    )
    await state.set_state(TranslationState.waiting_for_back_photo)

@router.callback_query(F.data == "skip_back_photo", TranslationState.waiting_for_back_photo)
async def skip_back_photo_callback(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await ask_for_doc_type(callback.message, state, edit_message=True)

async def ask_for_doc_type(message: types.Message, state: FSMContext, edit_message: bool = False):
    """Отправляет или обновляет сообщение для выбора типа документа."""
    buttons = []
    doc_types = doc_manager.get_types()
    for doc_id, doc_info in doc_types.items():
        buttons.append([
            InlineKeyboardButton(
                text=f"{doc_info['emoji']} {doc_info['name']}", 
                callback_data=f"doctype_{doc_id}"
            )
        ])
        
    keyboard = InlineKeyboardMarkup(inline_keyboard=buttons)
    text = "📸 Все фото получены!\nПожалуйста, выберите тип документа, чтобы мы могли правильно его перевести:"

    if edit_message:
        await message.edit_text(text, reply_markup=keyboard)
    else:
        await message.reply(text, reply_markup=keyboard)
        
    await state.set_state(TranslationState.choosing_doc_type)

@router.callback_query(F.data.startswith("doctype_"), TranslationState.choosing_doc_type)
async def process_document_type(callback: CallbackQuery, state: FSMContext):
    """Обрабатывает выбор типа документа и предлагает выбрать язык перевода."""
    doc_type = callback.data.replace("doctype_", "")
    await state.update_data(doc_type=doc_type)
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇷🇺 Русский язык", callback_data="lang_ru")],
        [InlineKeyboardButton(text="🇬🇧 Английский язык", callback_data="lang_en")]
    ])
    
    await callback.answer()
    await callback.message.edit_text(
        "✅ Тип документа выбран.\n\n🌐 Теперь, пожалуйста, выберите язык, на который нужно перевести данные:",
        reply_markup=keyboard
    )
    await state.set_state(TranslationState.choosing_language)

async def download_photos(bot, user_id: int, file_ids: list, temp_dir: Path) -> list[Path]:
    """Скачивает фото во временную директорию."""
    photo_paths = []
    for i, file_id in enumerate(file_ids):
        photo_path = temp_dir / f"doc_{user_id}_{file_id}_{i}.jpg"
        file = await bot.get_file(file_id)
        await bot.download_file(file.file_path, destination=photo_path)
        photo_paths.append(photo_path)
    return photo_paths

@router.callback_query(F.data.startswith("lang_"), TranslationState.choosing_language)
async def process_language(callback: CallbackQuery, state: FSMContext):
    """Обрабатывает выбор языка и запускает пайплайн распознавания и перевода."""
    lang = callback.data.replace("lang_", "")
    data = await state.get_data()
    doc_type = data.get("doc_type")
    file_ids = data.get("file_ids", [])
    await state.clear()
    
    if not file_ids or not doc_type:
        await callback.message.edit_text("❌ Ошибка: сессия устарела или данные не найдены. Пожалуйста, отправьте фото документа заново.")
        return

    lang_name = "Русский" if lang == "ru" else "Английский"
    current_config = doc_manager.get_document_config(doc_type, lang)
    if not current_config:
        await callback.answer("Неизвестный тип документа", show_alert=True)
        return

    await callback.answer()
    processing_msg = await callback.message.edit_text(
        f"✅ Документ: **{current_config['name']}**\n✅ Язык перевода: **{lang_name}**\n\n⏳ Подготавливаю фото и отправляю на анализ ИИ...", 
        parse_mode="Markdown"
    )

    gemini_service = get_gemini_service()
    if not gemini_service:
        await processing_msg.edit_text("❌ Сервис перевода (Gemini) в данный момент недоступен. Проверьте API ключ.")
        return

    temp_dir = Path("temp")
    temp_dir.mkdir(exist_ok=True)
    photo_paths = []
    
    try:
        # 1. Скачивание файлов
        try:
            photo_paths = await download_photos(callback.bot, callback.from_user.id, file_ids, temp_dir)
        except Exception as e:
            await processing_msg.edit_text(f"❌ Не удалось скачать фото:\n{str(e)}\n\nПопробуйте отправить заново.")
            return

        # 2. Подготовка шаблона в отдельном потоке (I/O)
        template_path = Path("templates") / current_config["template"]
        output_path = temp_dir / f"result_{callback.from_user.id}_{file_ids[0]}.docx"
        await asyncio.to_thread(_create_temp_template, template_path, current_config["name"], lang_name)
        
        # 3. Распознавание через Gemini
        total_photos = len(photo_paths)
        await processing_msg.edit_text(f"⏳ Изображения ({total_photos} шт) анализируются ИИ Gemini...\n\nТип: {current_config['name']}\nПеревод на: {lang_name}")
        
        extracted_data = await gemini_service.extract_data_from_image(
            image_path=photo_paths,
            prompt=current_config["prompt"]
        )
        
        if "error" in extracted_data:
            raise ValueError(f"Ошибка Gemini: {extracted_data['error']}\n{extracted_data.get('raw_text', '')}")

        # 4. Обновление состояния и переход к валидации
        await state.update_data(
            extracted_data=extracted_data,
            template_path=str(template_path),
            output_path=str(output_path),
            lang_name=lang_name,
            lang_code=lang
        )
        await state.set_state(TranslationState.validating_data)
        
        await processing_msg.delete()
        await send_validation_menu(callback.message, extracted_data, lang_name, lang)
        
    except Exception as e:
        await processing_msg.edit_text(f"❌ Произошла ошибка при обработке:\n{str(e)}")
    finally:
        # Гарантированное удаление фото
        for p in photo_paths:
            if p.exists():
                try:
                    os.remove(p)
                except Exception:
                    pass

def get_validation_keyboard(data_dict: dict, lang_code: str = "ru") -> InlineKeyboardMarkup:
    buttons = []
    current_row = []
    
    for key in data_dict.keys():
        localized_name = doc_manager.localize_field(key, lang_code)
        btn_text = f"✏️ {localized_name[:20]}"
        cb_data = f"editf_{key}"
        if len(cb_data) > 64:
            cb_data = cb_data[:64]
            
        current_row.append(InlineKeyboardButton(text=btn_text, callback_data=cb_data))
        
        if len(current_row) == 2:
            buttons.append(current_row)
            current_row = []
            
    if current_row:
        buttons.append(current_row)
    
    confirm_text = "✅ Подтвердить и создать" if lang_code == "ru" else "✅ Confirm and generate"
    buttons.append([InlineKeyboardButton(text=confirm_text, callback_data="confirm_generation")])
    return InlineKeyboardMarkup(inline_keyboard=buttons)

async def send_validation_menu(message: types.Message, data_dict: dict, lang_name: str = "выбранный", lang_code: str = "ru"):
    title = f"🤖 **Результат распознавания ({lang_name}):**\n" if lang_code == "ru" else f"🤖 **Extraction Result ({lang_name}):**\n"
    lines = [title]
    for k, v in data_dict.items():
        localized_name = doc_manager.localize_field(k, lang_code)
        lines.append(f"▪️ **{localized_name}**: `{v}`")
        
    instr = "\n_Нажмите на кнопку с именем поля ниже, если нужно его исправить._" if lang_code == "ru" else "\n_Click on the field name below if you need to manually fix it._"
    text = "\n".join(lines) + "\n" + instr
    
    await message.answer(text, reply_markup=get_validation_keyboard(data_dict, lang_code), parse_mode="Markdown")

@router.callback_query(F.data.startswith("editf_"), TranslationState.validating_data)
async def process_edit_field_selection(callback: CallbackQuery, state: FSMContext):
    field_key = callback.data.replace("editf_", "")
    await state.update_data(editing_field_name=field_key)
    await state.set_state(TranslationState.editing_field)
    
    await callback.answer()
    
    data = await state.get_data()
    extracted_data = data.get("extracted_data", {})
    current_value = extracted_data.get(field_key, "")
    lang_code = data.get("lang_code", "ru")
    localized_name = doc_manager.localize_field(field_key, lang_code)
    
    prompt = (
        f"✏️ Введите новое значение для поля **{localized_name}**\nТекущее значение: `{current_value}`\n\n(Или отправьте /cancel для отмены редактирования)" 
        if lang_code == "ru" 
        else f"✏️ Enter new value for **{localized_name}**\nCurrent value: `{current_value}`\n\n(Or send /cancel to abort)"
    )
    
    await callback.message.answer(prompt, parse_mode="Markdown")

@router.message(TranslationState.editing_field)
async def process_new_field_value(message: types.Message, state: FSMContext):
    data = await state.get_data()
    extracted_data = data.get("extracted_data", {})
    lang_code = data.get("lang_code", "ru")
    lang_name = data.get("lang_name", "выбранный")

    if message.text == "/cancel":
        await state.set_state(TranslationState.validating_data)
        await send_validation_menu(message, extracted_data, lang_name, lang_code)
        return

    new_value = message.text
    field_key = data.get("editing_field_name")
    
    if field_key:
        extracted_data[field_key] = new_value
        await state.update_data(extracted_data=extracted_data)
    
    await state.set_state(TranslationState.validating_data)
    ok_msg = "✅ Значение изменено!" if lang_code == "ru" else "✅ Value updated!"
    await message.reply(ok_msg)
    await send_validation_menu(message, extracted_data, lang_name, lang_code)

@router.callback_query(F.data == "confirm_generation", TranslationState.validating_data)
async def confirm_generation(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    processing_msg = await callback.message.edit_text("⏳ Генерирую документ...")
    
    data = await state.get_data()
    extracted_data = data.get("extracted_data", {})
    
    template_path_str = data.get("template_path")
    output_path_str = data.get("output_path")
    
    if not template_path_str or not output_path_str:
        await processing_msg.edit_text("❌ Ошибка: пути к файлам не найдены в сессии. Начните заново.")
        await state.clear()
        return
        
    template_path = Path(template_path_str)
    output_path = Path(output_path_str)
    lang_name = data.get("lang_name", "Выбранный")
    
    gemini_service = get_gemini_service()
    if not gemini_service:
        await processing_msg.edit_text("❌ Сервис перевода (Gemini) в данный момент недоступен.")
        return

    try:
        # Вставляем данные в Word (I/O, в отдельном потоке)
        await asyncio.to_thread(_generate_docx, gemini_service, extracted_data, template_path, output_path)
        
        # Отправляем готовый Word
        await processing_msg.edit_text("✅ Документ готов! Отправляю файл...")
        document = FSInputFile(output_path)
        await callback.message.answer_document(document, caption=f"Вот ваш проверенный перевод на {lang_name.lower()} язык!")
        
    except Exception as e:
        await processing_msg.edit_text(f"❌ Произошла ошибка при генерации документа:\n{str(e)}")
        
    finally:
        if output_path.exists():
            try:
                os.remove(output_path)
            except Exception:
                pass
        await state.clear()
