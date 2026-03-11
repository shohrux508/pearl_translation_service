import logging
from pathlib import Path
from typing import Dict, Any, Union
from docx import Document
from docxtpl import DocxTemplate

import re
logger = logging.getLogger(__name__)

def _sanitize_for_xml(text: str) -> str:
    """Удаляет control-символы, которые ломают XML в MS Word."""
    if not isinstance(text, str):
        return text
    # Удаляем 0x00-0x08, 0x0B-0x0C, 0x0E-0x1F, 0x7F-0x9F
    cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    # Заменяем обычные переносы строк на пробелы, так как они могут
    # вызывать неопознанную ошибку парсинга внутри тегов <w:t> в Word
    cleaned = cleaned.replace('\r\n', ' ').replace('\n', ' ')
    return cleaned

def _sanitize_data(data: Any) -> Any:
    if isinstance(data, dict):
        return {k: _sanitize_data(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [_sanitize_data(v) for v in data]
    elif isinstance(data, str):
        return _sanitize_for_xml(data)
    elif data is None:
        return ""
    return data

class DocxService:
    """
    Сервис для работы с Word-документами (.docx).
    Отвечает за генерацию и заполнение шаблонов.
    """
    
    def create_temp_template(self, template_path: Path, doc_name: str, lang_name: str) -> None:
        """
        Создает пустой тестовый шаблон, если его не существует.
        """
        if not template_path.exists():
            template_path.parent.mkdir(exist_ok=True, parents=True)
            doc = Document()
            doc.add_paragraph(f"Временный тестовый шаблон для: {doc_name}")
            doc.add_paragraph(f"Язык шаблона (и перевода): {lang_name}")
            doc.add_paragraph("Данные будут автоматически вставляться, если шаблонизатор найдет совпадающие поля (например {{surname}}).")
            doc.save(template_path)
            logger.info(f"Создан временный шаблон: {template_path}")

    def generate_docx(
        self, 
        data: Dict[str, Any], 
        template_path: Union[str, Path], 
        output_path: Union[str, Path]
    ) -> str:
        """
        Берет JSON-словарь извлеченных данных и заменяет плейсхолдеры 
        (например, {{key}}) в .docx шаблоне.
        """
        logger.info(f"Открытие шаблона {template_path} для вставки данных...")
        doc = DocxTemplate(template_path)
        
        # Санитизируем данные: Word очень строг к непечатаемым символам в XML
        safe_data = _sanitize_data(data)
        
        doc.render(safe_data)
        doc.save(output_path)
        logger.info(f"Документ успешно сохранен: {output_path}")
        return output_path
